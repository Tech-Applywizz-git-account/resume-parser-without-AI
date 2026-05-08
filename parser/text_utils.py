# text_utils.py

import os
import pdfplumber
import fitz  # PyMuPDF
import docx
import re

def extract_links_from_pdf(file_path):
    links = []
    try:
        doc = fitz.open(file_path)
        for page in doc:
            for link in page.get_links():
                if 'uri' in link:
                    links.append(link['uri'])
    except Exception as e:
        print(f"Error extracting PDF links: {e}")
    return list(set(links))

def extract_links_from_docx(file_path):
    links = []
    try:
        doc = docx.Document(file_path)
        rels = doc.part.rels
        for rel in rels.values():
            if "hyperlink" in rel.reltype:
                links.append(rel._target)
    except Exception as e:
        print(f"Error extracting DOCX links: {e}")
    return list(set(links))

def extract_text_from_pdf_plumber(file_path):
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t: text += t + "\n"
    return text

def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    links = []
    
    if ext == '.pdf':
        text = extract_text_from_pdf_plumber(file_path)
        links = extract_links_from_pdf(file_path)
    elif ext == '.docx':
        doc = docx.Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        links = extract_links_from_docx(file_path)
    else:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
            links = re.findall(r'https?://[^\s]+', text)
            
    return text, list(set(links))

def clean_text(text):
    """Clean extra spaces but PRESERVE newlines for structural analysis"""
    if not text: return ""
    # Remove excessive horizontal spaces but keep \n
    text = re.sub(r'[ \t]+', ' ', text)
    # Remove excessive vertical spaces
    text = re.sub(r'\n\s*\n', '\n\n', text)
    return text.strip()












# """
# text_utils.py — File I/O, text extraction, and structural normalization.

# Veteran notes
# ─────────────
# • Two-column PDFs are the #1 silent killer of naive parsers.  pdfplumber
#   returns words with (x0, y0) coordinates; we use those to reconstruct
#   reading order for both 1- and 2-column layouts.

# • PyMuPDF (fitz) is better for link/annotation extraction but pdfplumber is
#   better for text with bounding boxes.  We use both.

# • Some PDFs encode characters as ligatures (ﬁ, ﬂ, ﬀ) — normalise them.

# • DOCX files may embed hyperlinks only in the relationships part, not in
#   paragraph runs — check both places.

# • Always fall back to raw text extraction when structured parsing fails.
# """

# import os
# import re
# import unicodedata
# from collections import defaultdict

# import pdfplumber
# import fitz          # PyMuPDF
# import docx

# # ──────────────────────────────────────────────────────────────────────────────
# # Unicode / encoding normalisation
# # ──────────────────────────────────────────────────────────────────────────────

# _LIGATURE_MAP = str.maketrans({
#     '\ufb00': 'ff', '\ufb01': 'fi', '\ufb02': 'fl',
#     '\ufb03': 'ffi', '\ufb04': 'ffl', '\ufb05': 'st',
#     '\u2019': "'",  '\u2018': "'",  '\u201c': '"',  '\u201d': '"',
#     '\u2013': '-',  '\u2014': '-',  '\u2022': '•',  '\u00a0': ' ',
#     '\u200b': '',   '\ufeff': '',
# })

# def _fix_unicode(text: str) -> str:
#     """Normalise ligatures, curly quotes, en-/em-dashes, NBSP, zero-width."""
#     text = unicodedata.normalize('NFKC', text)
#     return text.translate(_LIGATURE_MAP)


# # ──────────────────────────────────────────────────────────────────────────────
# # PDF helpers
# # ──────────────────────────────────────────────────────────────────────────────

# _TWO_COL_GAP = 0.35   # if a page's words span two clusters > 35 % of page width apart → two columns


# def _detect_columns(words: list[dict], page_width: float) -> list[list[dict]]:
#     """
#     Cluster words into columns by their x0 position.

#     Returns a list of word groups (columns), left to right.
#     """
#     if not words:
#         return [words]

#     # Build x0 histogram; look for a gap in the middle third of the page
#     mid_lo = page_width * 0.30
#     mid_hi = page_width * 0.70

#     left_words  = [w for w in words if w['x1'] < mid_hi]
#     right_words = [w for w in words if w['x0'] > mid_lo]

#     # If there's significant overlap there's only one column
#     if len(left_words) > 0 and len(right_words) > 0:
#         overlap = [w for w in words if w['x0'] > mid_lo and w['x1'] < mid_hi]
#         col_gap_ratio = 1 - len(overlap) / len(words)
#         if col_gap_ratio > _TWO_COL_GAP:
#             col_split = (mid_lo + mid_hi) / 2
#             left  = [w for w in words if w['x0'] < col_split]
#             right = [w for w in words if w['x0'] >= col_split]
#             return [left, right]

#     return [words]


# def _words_to_text(words: list[dict]) -> str:
#     """
#     Convert a pdfplumber word list (with x0, y0, text) into a readable string,
#     respecting line breaks based on y-coordinate proximity.
#     """
#     if not words:
#         return ""

#     # Sort by (row, x)
#     Y_TOLERANCE = 4   # pixels; words within this y-distance are on the same line
#     rows: dict[int, list[dict]] = defaultdict(list)
#     for w in words:
#         y_key = round(w['top'] / Y_TOLERANCE) * Y_TOLERANCE
#         rows[y_key].append(w)

#     lines = []
#     for y_key in sorted(rows):
#         row_words = sorted(rows[y_key], key=lambda w: w['x0'])
#         line = ' '.join(w['text'] for w in row_words)
#         lines.append(line)

#     return '\n'.join(lines)


# def extract_text_from_pdf(file_path: str) -> str:
#     """
#     Extract text from a PDF, handling:
#       • single-column layouts
#       • two-column layouts
#       • scanned PDFs (fallback to PyMuPDF)
#     """
#     full_text = ""

#     try:
#         with pdfplumber.open(file_path) as pdf:
#             for page in pdf.pages:
#                 words = page.extract_words(
#                     x_tolerance=3, y_tolerance=3,
#                     keep_blank_chars=False, use_text_flow=False,
#                     extra_attrs=['fontname', 'size'],
#                 )

#                 if not words:
#                     # Empty page or scanned image — try basic text extraction
#                     t = page.extract_text()
#                     if t:
#                         full_text += t + "\n"
#                     continue

#                 columns = _detect_columns(words, page.width)
#                 for col in columns:
#                     col_text = _words_to_text(col)
#                     full_text += col_text + "\n"

#                 full_text += "\n"  # page separator

#     except Exception as exc:
#         print(f"[text_utils] pdfplumber failed: {exc}; trying PyMuPDF fallback")
#         full_text = _fitz_fallback(file_path)

#     return _fix_unicode(full_text)


# def _fitz_fallback(file_path: str) -> str:
#     """Use PyMuPDF for text extraction when pdfplumber fails."""
#     text = ""
#     try:
#         doc = fitz.open(file_path)
#         for page in doc:
#             text += page.get_text("text") + "\n"
#     except Exception as exc:
#         print(f"[text_utils] PyMuPDF fallback also failed: {exc}")
#     return text


# def extract_links_from_pdf(file_path: str) -> list[str]:
#     """Extract all hyperlinks from PDF annotations."""
#     links = []
#     try:
#         doc = fitz.open(file_path)
#         for page in doc:
#             for link in page.get_links():
#                 uri = link.get('uri', '')
#                 if uri:
#                     links.append(uri)
#     except Exception as exc:
#         print(f"[text_utils] PDF link extraction error: {exc}")
#     return list(set(links))


# # ──────────────────────────────────────────────────────────────────────────────
# # DOCX helpers
# # ──────────────────────────────────────────────────────────────────────────────

# def extract_text_from_docx(file_path: str) -> str:
#     """
#     Extract DOCX text preserving paragraph structure.
#     Also handles text inside tables (very common in resumes).
#     """
#     parts = []
#     try:
#         doc = docx.Document(file_path)

#         # Regular paragraphs
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 parts.append(para.text)

#         # Table cells
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     for para in cell.paragraphs:
#                         if para.text.strip():
#                             parts.append(para.text)

#     except Exception as exc:
#         print(f"[text_utils] DOCX extraction error: {exc}")

#     return _fix_unicode('\n'.join(parts))


# def extract_links_from_docx(file_path: str) -> list[str]:
#     """Extract hyperlinks from DOCX relationship part."""
#     links = []
#     try:
#         doc = docx.Document(file_path)
#         for rel in doc.part.rels.values():
#             if "hyperlink" in rel.reltype:
#                 target = str(rel._target)
#                 if target.startswith('http'):
#                     links.append(target)
#     except Exception as exc:
#         print(f"[text_utils] DOCX link extraction error: {exc}")
#     return list(set(links))


# # ──────────────────────────────────────────────────────────────────────────────
# # Dispatcher
# # ──────────────────────────────────────────────────────────────────────────────

# def extract_text_from_file(file_path: str) -> tuple[str, list[str]]:
#     """
#     Main entry point.  Returns (text, links).
#     Supports .pdf, .docx, .doc, .txt, .rtf.
#     """
#     ext = os.path.splitext(file_path)[1].lower()

#     if ext == '.pdf':
#         text  = extract_text_from_pdf(file_path)
#         links = extract_links_from_pdf(file_path)

#     elif ext in ('.docx', '.doc'):
#         text  = extract_text_from_docx(file_path)
#         links = extract_links_from_docx(file_path)

#     else:
#         # Plain text / RTF / fallback
#         try:
#             with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
#                 text = f.read()
#         except Exception as exc:
#             print(f"[text_utils] Plain text read error: {exc}")
#             text = ""
#         links = re.findall(r'https?://[^\s]+', text)

#     # Always also scrape any inline URLs from the text itself
#     inline_links = re.findall(r'https?://[^\s<>"\'\\)]{4,}', text)
#     all_links    = list(set(links + inline_links))

#     return text, all_links


# # ──────────────────────────────────────────────────────────────────────────────
# # Text cleaning
# # ──────────────────────────────────────────────────────────────────────────────

# def clean_text(text: str) -> str:
#     """
#     Light cleaning that preserves structure:
#     • collapse multiple spaces/tabs on a single line
#     • collapse 3+ blank lines to 2
#     • strip trailing whitespace per line
#     Does NOT remove newlines — the segmenter depends on them.
#     """
#     if not text:
#         return ""

#     # Per-line: strip trailing spaces, collapse internal horizontal whitespace
#     lines = [re.sub(r'[ \t]+', ' ', line).rstrip() for line in text.split('\n')]
#     text  = '\n'.join(lines)

#     # Collapse excessive blank lines
#     text  = re.sub(r'\n{3,}', '\n\n', text)

#     return text.strip()